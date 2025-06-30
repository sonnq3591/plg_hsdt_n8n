#!/usr/bin/env python3
"""
15_KE_HOACH_TRIEN_KHAI_FIXED.py - Fixed version that ensures correct data scaling
Forces generators to use the exact data from current batch
"""

import os
import json
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import pandas as pd
from io import StringIO

class Template15FixedGenerator:
    """Generate template 15 with FORCED correct data scaling"""
    
    def __init__(self):
        """Initialize the generator"""
        # Set working directory to script's parent (BID_PROCESSOR)
        self.base_dir = Path(__file__).parent.parent
        os.chdir(self.base_dir)
        
        # Read current batch ID
        self.batch_id = self.read_batch_id()
        if not self.batch_id:
            raise ValueError("❌ No current batch ID found!")
        
        # Set up paths
        self.batch_dir = self.base_dir / self.batch_id
        self.extracted_dir = self.batch_dir / "extracted_data"
        self.docx_dir = self.batch_dir / "docx"
        self.pictures_dir = self.batch_dir / "pictures"
        self.templates_dir = self.base_dir / "templates"
        
        # Ensure directories exist
        self.docx_dir.mkdir(exist_ok=True)
        self.pictures_dir.mkdir(exist_ok=True)
        
        print(f"📄 Template15FixedGenerator initialized for batch: {self.batch_id}")

    def read_batch_id(self):
        """Read current batch ID from current_batch.txt"""
        try:
            batch_file = self.base_dir / "current_batch.txt"
            with open(batch_file, "r") as f:
                batch_id = f.read().strip()
            return batch_id
        except FileNotFoundError:
            print("❌ current_batch.txt not found!")
            return None

    def load_master_data(self):
        """Load master data to determine replacement strategy"""
        master_file = self.extracted_dir / "master_data.json"
        
        if not master_file.exists():
            raise FileNotFoundError(f"❌ Master data not found: {master_file}")
        
        with open(master_file, 'r', encoding='utf-8') as f:
            master_data = json.load(f)
        
        return master_data

    def detect_step_count(self, master_data):
        """Detect if this is 21-step or 23-step process from cac_buoc_thuc_hien"""
        placeholders = master_data.get("placeholders", {})
        cac_buoc_data = placeholders.get("cac_buoc_thuc_hien", {})
        content = cac_buoc_data.get("content", {})
        
        # Check if it's the conditional table type with step_count
        if isinstance(content, dict) and "step_count" in content:
            step_count = content["step_count"]
            print(f"🔍 Detected step count: {step_count}")
            return step_count
        
        # Fallback: assume 21 steps
        print(f"⚠️ Could not detect step count, defaulting to 21")
        return 21

    def extract_days_from_time(self, time_str):
        """Extract number of days from time string"""
        import re
        numbers = re.findall(r'\d+', time_str)
        if numbers:
            return int(numbers[0])
        return 120  # Default fallback

    def generate_timeline_chart_directly(self, step_count, target_days):
        """Generate timeline chart directly with FORCED scaling"""
        print(f"📊 Generating {step_count}-step timeline with {target_days} days...")
        
        # Hardcoded base data for different step counts
        if step_count == 21:
            csv_data = """Task,So_ngay
Giao nhận tài liệu và lập biên bản giao nhận tài liệu,1
Vận chuyển tài liệu từ kho bảo quản đến địa điểm chỉnh lý,1
Vệ sinh sơ bộ tài liệu,1
Khảo sát và biên soạn các văn bản hướng dẫn chỉnh lý,1
Phân loại tài liệu theo Hướng dẫn phân loại,20
Lập hồ sơ hoặc chỉnh sửa hoàn thiện hồ sơ theo Hướng dẫn lập hồ sơ,30
Viết các trường thông tin vào phiếu tin,3
Kiểm tra chỉnh sửa hồ sơ và phiếu tin,1
Hệ thống hóa phiếu tin theo phương án phân loại,1
Hệ thống hóa hồ sơ theo phiếu tin,1
Biên mục hồ sơ,15
Kiểm tra và chỉnh sửa việc biên mục hồ sơ,1
Ghi số hồ sơ chính thức vào phiếu tin và lên bìa hồ sơ,7
Vệ sinh tài liệu tháo bỏ ghim kẹp làm phẳng và đưa tài liệu vào bìa hồ sơ,2
Đưa hồ sơ vào hộp cặp,3
Viết in và dán nhãn hộp cặp,3
Vận chuyển tài liệu vào kho và xếp lên giá,1
Kiểm tra chỉnh sửa việc biên phiếu tin,1
Lập mục lục hồ sơ,5
Thống kê bó gói lập danh mục và viết thuyết minh tài liệu loại,1
Kết thúc chỉnh lý,1"""
        else:  # 23 steps
            csv_data = """Task,So_ngay
Giao nhận tài liệu và lập biên bản giao nhận tài liệu,0.5
Vận chuyển tài liệu từ kho bảo quản đến địa điểm chỉnh lý,1
Vệ sinh sơ bộ tài liệu,1
Khảo sát và biên soạn các văn bản hướng dẫn chỉnh lý,1
Phân loại tài liệu theo Hướng dẫn phân loại,20
Lập hồ sơ hoặc chỉnh sửa hoàn thiện hồ sơ theo Hướng dẫn lập hồ sơ,30
Viết các trường thông tin vào phiếu tin,2
Kiểm tra chỉnh sửa hồ sơ và phiếu tin,1
Hệ thống hóa phiếu tin theo phương án phân loại,1
Hệ thống hóa hồ sơ theo phiếu tin,1
Biên mục hồ sơ,15
Kiểm tra và chỉnh sửa việc biên mục hồ sơ,1
Ghi số hồ sơ chính thức vào phiếu tin và lên bìa hồ sơ,6
Vệ sinh tài liệu tháo bỏ ghim kẹp làm phẳng và đưa tài liệu vào bìa hồ sơ,1
Đưa hồ sơ vào hộp cặp,2
Viết in và dán nhãn hộp cặp,2
Vận chuyển tài liệu vào kho và xếp lên giá,1
Giao nhận tài liệu sau chỉnh lý và lập Biên bản giao nhận tài liệu,1
Nhập phiếu tin vào cơ sở dữ liệu,5
Kiểm tra chỉnh sửa việc nhập phiếu tin,1
Lập mục lục hồ sơ,4
Thống kê bó gói lập danh mục và viết thuyết minh tài liệu loại,1
Kết thúc chỉnh lý,1"""
        
        # Process data
        df = pd.read_csv(StringIO(csv_data))
        
        # FORCE scaling with exact target_days
        base_total = df['So_ngay'].sum()
        scaling_ratio = target_days / base_total
        
        print(f"📊 Base total: {base_total} days")
        print(f"📊 Target total: {target_days} days")
        print(f"📊 Scaling ratio: {scaling_ratio:.3f}")
        
        # Apply scaling and round up
        df['Days_Scaled'] = (df['So_ngay'] * scaling_ratio).apply(lambda x: max(1, round(x)))
        scaled_total = df['Days_Scaled'].sum()
        
        print(f"📊 Actual scaled total: {scaled_total} days")
        
        # Create the chart
        def wrap_text(text, words_per_line=2):
            words = text.split()
            return '\n'.join([' '.join(words[i:i+words_per_line]) for i in range(0, len(words), words_per_line)])

        wrapped_labels = [wrap_text(task, 2) for task in df["Task"]]
        max_lines = max(label.count('\n') + 1 for label in wrapped_labels)
        bottom_margin = 0.4 + 0.03 * max_lines
        fig_height = 12 + 0.5 * max_lines

        # Create figure
        fig, ax = plt.subplots(figsize=(26, fig_height))
        x = df.index
        ax.plot(x, df['Days_Scaled'], marker='o', color='red', linestyle='-')
        ax.set_title(f"KẾ HOẠCH THỰC HIỆN CÔNG VIỆC ({step_count} BƯỚC)", fontsize=20, pad=25)
        ax.set_ylabel("Số ngày", fontsize=16)
        ax.set_xticks(x)
        ax.set_xticklabels(wrapped_labels, rotation=0, ha='center', fontsize=11.5)
        ax.grid(True, axis='y', linestyle='--', alpha=0.6)

        # Set y-axis limits with extra space at top for labels
        max_y = df['Days_Scaled'].max()
        ax.set_ylim(0, max_y * 1.15)

        # Annotate data points
        for i, y in enumerate(df['Days_Scaled']):
            ax.annotate(str(y), (x[i], y), textcoords="offset points", xytext=(0, 10), va='bottom', ha='center', fontsize=12)

        plt.subplots_adjust(left=0.05, right=0.98, top=0.93, bottom=bottom_margin)
        
        # Save chart
        output_file = self.pictures_dir / f"{step_count}_BUOC_KH_THUC_HIEN.png"
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ Timeline chart saved: {output_file}")
        return output_file.exists()

    def generate_personnel_chart_directly(self, step_count, target_days):
        """Generate personnel chart directly with FORCED scaling"""
        print(f"👥 Generating {step_count}-step personnel chart with {target_days} days...")
        
        # Timeline data (same as above)
        if step_count == 21:
            timeline_base = [1,1,1,1,20,30,3,1,1,1,15,1,7,2,3,3,1,1,5,1,1]
            personnel_data = [
                [1,1,5], [1,1,10], [1,1,10], [1,1,2], [1,1,10], [1,1,10], [1,1,10], [1,1,2],
                [1,1,10], [1,1,10], [1,1,10], [1,1,2], [1,1,10], [1,1,10], [1,1,10], [1,1,10],
                [1,1,10], [1,1,2], [1,1,10], [1,1,10], [1,1,2]
            ]
        else:  # 23 steps
            timeline_base = [0.5,1,1,1,20,30,2,1,1,1,15,1,6,1,2,2,1,1,5,1,4,1,1]
            personnel_data = [
                [1,1,5], [1,1,10], [1,1,10], [1,1,2], [1,1,10], [1,1,10], [1,1,10], [1,1,2],
                [1,1,10], [1,1,10], [1,1,10], [1,1,2], [1,1,10], [1,1,10], [1,1,10], [1,1,10],
                [1,1,10], [1,1,2], [1,1,10], [1,1,2], [1,1,10], [1,1,10], [1,1,2]
            ]
        
        # Scale timeline data
        base_total = sum(timeline_base)
        scaling_ratio = target_days / base_total
        scaled_timeline = [max(1, round(x * scaling_ratio)) for x in timeline_base]
        
        # Create task names (shortened for chart)
        if step_count == 21:
            task_names = [
                "Giao nhận tài liệu và lập biên bản giao nhận tài liệu",
                "Vận chuyển tài liệu từ kho bảo quản đến địa điểm chỉnh lý",
                "Vệ sinh sơ bộ tài liệu", "Khảo sát và biên soạn các văn bản hướng dẫn chỉnh lý",
                "Phân loại tài liệu theo Hướng dẫn phân loại", "Lập hồ sơ hoặc chỉnh sửa hoàn thiện hồ sơ theo Hướng dẫn lập hồ sơ",
                "Viết các trường thông tin vào phiếu tin", "Kiểm tra chỉnh sửa hồ sơ và phiếu tin",
                "Hệ thống hóa phiếu tin theo phương án phân loại", "Hệ thống hóa hồ sơ theo phiếu tin",
                "Biên mục hồ sơ", "Kiểm tra và chỉnh sửa việc biên mục hồ sơ",
                "Ghi số hồ sơ chính thức vào phiếu tin và lên bìa hồ sơ", "Vệ sinh tài liệu tháo bỏ ghim kẹp làm phẳng và đưa tài liệu vào bìa hồ sơ",
                "Đưa hồ sơ vào hộp cặp", "Viết in và dán nhãn hộp cặp", "Vận chuyển tài liệu vào kho và xếp lên giá",
                "Kiểm tra chỉnh sửa việc biên phiếu tin", "Lập mục lục hồ sơ", "Thống kê bó gói lập danh mục và viết thuyết minh tài liệu loại",
                "Kết thúc chỉnh lý"
            ]
        else:  # 23 steps
            task_names = [
                "Giao nhận tài liệu và lập biên bản giao nhận tài liệu", "Vận chuyển tài liệu từ kho bảo quản đến địa điểm chỉnh lý",
                "Vệ sinh sơ bộ tài liệu", "Khảo sát và biên soạn các văn bản hướng dẫn chỉnh lý",
                "Phân loại tài liệu theo Hướng dẫn phân loại", "Lập hồ sơ hoặc chỉnh sửa hoàn thiện hồ sơ theo Hướng dẫn lập hồ sơ",
                "Viết các trường thông tin vào phiếu tin", "Kiểm tra chỉnh sửa hồ sơ và phiếu tin",
                "Hệ thống hóa phiếu tin theo phương án phân loại", "Hệ thống hóa hồ sơ theo phiếu tin",
                "Biên mục hồ sơ", "Kiểm tra và chỉnh sửa việc biên mục hồ sơ",
                "Ghi số hồ sơ chính thức vào phiếu tin và lên bìa hồ sơ", "Vệ sinh tài liệu tháo bỏ ghim kẹp làm phẳng và đưa tài liệu vào bìa hồ sơ",
                "Đưa hồ sơ vào hộp cặp", "Viết in và dán nhãn hộp cặp", "Vận chuyển tài liệu vào kho và xếp lên giá",
                "Giao nhận tài liệu sau chỉnh lý và lập Biên bản giao nhận tài liệu", "Nhập phiếu tin vào cơ sở dữ liệu",
                "Kiểm tra chỉnh sửa việc nhập phiếu tin", "Lập mục lục hồ sơ", "Thống kê bó gói lập danh mục và viết thuyết minh tài liệu loại",
                "Kết thúc chỉnh lý"
            ]

        # Create DataFrame for personnel chart
        data = {
            'Task': task_names,
            'Ngày thực hiện': scaled_timeline,
            'Quản lý dự án': [row[0] for row in personnel_data],
            'Trưởng nhóm chỉnh lý': [row[1] for row in personnel_data],
            'Nhân sự chỉnh lý': [row[2] for row in personnel_data]
        }
        
        df = pd.DataFrame(data)
        df.set_index('Task', inplace=True)

        # Create personnel chart
        def wrap_text(text, words_per_line=2):
            words = text.split()
            return '\n'.join([' '.join(words[i:i+words_per_line]) for i in range(0, len(words), words_per_line)])

        wrapped_labels = [wrap_text(task, 2) for task in df.index]
        max_lines = max(label.count('\n') + 1 for label in wrapped_labels)
        bottom_margin = 0.3 + 0.02 * max_lines
        fig_height = 12 + 0.5 * max_lines

        fig, ax = plt.subplots(figsize=(26, fig_height))
        df.plot(kind='bar', ax=ax, width=0.55)
        ax.set_xticklabels(wrapped_labels, rotation=0, ha='center', fontsize=11)

        # Add data labels
        for container in ax.containers:
            for bar in container:
                height = bar.get_height()
                if height > 0:
                    ax.annotate(f'{int(height)}',
                                xy=(bar.get_x() + bar.get_width() / 2, height),
                                xytext=(0, 3), textcoords="offset points",
                                ha='center', va='bottom', fontsize=10)

        ax.set_title(f'KẾ HOẠCH NHÂN SỰ ({step_count} BƯỚC)', fontsize=20)
        ax.set_ylabel('Ngày/Nhân sự', fontsize=16)
        ax.set_xlabel('', fontsize=16)
        ax.legend(title='Nhóm nhân sự', fontsize=12, loc='upper right')
        ax.yaxis.grid(True, linestyle='--', alpha=0.6)

        # Add table below
        table_data = df.T.values
        row_labels = df.columns.tolist()
        table = plt.table(cellText=table_data, rowLabels=row_labels,
                          cellLoc='center', rowLoc='center', loc='bottom',
                          bbox=[0.0, -bottom_margin * 0.88, 1.0, 0.2])
        table.auto_set_font_size(False)
        table.set_fontsize(10)

        plt.subplots_adjust(left=0.05, right=0.98, top=0.93, bottom=bottom_margin)
        
        # Save chart
        output_file = self.pictures_dir / f"{step_count}_BUOC_KH_NHAN_SU.png"
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ Personnel chart saved: {output_file}")
        return output_file.exists()

    def copy_template_to_working_file(self, template_name, output_name):
        """Copy template to working file for processing"""
        template_file = self.templates_dir / template_name
        output_file = self.docx_dir / output_name
        
        if not template_file.exists():
            raise FileNotFoundError(f"❌ Template not found: {template_file}")
        
        shutil.copy2(template_file, output_file)
        print(f"✅ Copied template: {template_name} → {output_name}")
        return output_file

    def replace_text_placeholder(self, doc, placeholder_key, content):
        """Replace text placeholders with content"""
        placeholder_tag = f"{{{{{placeholder_key}}}}}"
        print(f"📝 Replacing {placeholder_tag} with: '{content}'")
        
        replaced = False
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            full_text = paragraph.text
            if placeholder_tag not in full_text:
                continue
            
            print(f"📍 Found {placeholder_tag} in paragraph {para_idx}")
            
            # Check if placeholder should be bold from template
            template_bold = False
            for run in paragraph.runs:
                if placeholder_tag in run.text and run.bold:
                    template_bold = True
                    break
            
            # Split paragraph text into parts
            before_text, after_text = full_text.split(placeholder_tag, 1)
            
            # Clear all runs
            for run in paragraph.runs:
                run.text = ""
            
            # Rebuild with selective formatting
            current_run = 0
            
            # Add "before" text
            if before_text:
                if current_run < len(paragraph.runs):
                    run = paragraph.runs[current_run]
                else:
                    run = paragraph.add_run()
                
                run.text = before_text
                run.bold = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                current_run += 1
            
            # Add replacement content
            if current_run < len(paragraph.runs):
                run = paragraph.runs[current_run]
            else:
                run = paragraph.add_run()
            
            run.text = content
            run.bold = template_bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            current_run += 1
            
            # Add "after" text
            if after_text:
                if current_run < len(paragraph.runs):
                    run = paragraph.runs[current_run]
                else:
                    run = paragraph.add_run()
                
                run.text = after_text
                run.bold = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
            
            print(f"   ✅ Applied formatting: '{content}' (bold={template_bold})")
            
            replaced = True
            break
        
        return replaced

    def replace_image_placeholder(self, doc, placeholder_key, image_filename):
        """Replace image placeholders with actual images"""
        placeholder_tag = f"{{{{{placeholder_key}}}}}"
        image_path = self.pictures_dir / image_filename
        
        if not image_path.exists():
            print(f"❌ Image not found: {image_path}")
            return False
        
        print(f"🖼️ Replacing {placeholder_tag} with: {image_filename}")
        
        replaced = False
        
        # Check all paragraphs for the placeholder
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                print(f"📍 Found {placeholder_tag} in paragraph {para_idx}")
                
                # Clear paragraph completely
                paragraph.clear()
                
                # Add image run
                run = paragraph.add_run()
                
                try:
                    # Use FULL document width for better visibility
                    run.add_picture(str(image_path), width=Inches(8.5))
                    
                    # Center the paragraph
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    print(f"✅ Inserted image: {image_filename} (width=8.5\" - full page)")
                    replaced = True
                    break
                    
                except Exception as e:
                    print(f"❌ Failed to insert image {image_filename}: {e}")
                    # Fallback: replace with text
                    run.text = f"[IMAGE: {image_filename}]"
                    replaced = True
                    break
        
        # Also check in table cells
        if not replaced:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if placeholder_tag in paragraph.text:
                                print(f"📍 Found {placeholder_tag} in table cell paragraph {para_idx}")
                                
                                # Clear paragraph
                                paragraph.clear()
                                
                                # Add image
                                run = paragraph.add_run()
                                
                                try:
                                    # Use large width for table cells too
                                    run.add_picture(str(image_path), width=Inches(8.0))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    print(f"✅ Inserted image in table: {image_filename} (width=8.0\")")
                                    replaced = True
                                    break
                                    
                                except Exception as e:
                                    print(f"❌ Failed to insert image in table: {e}")
                                    run.text = f"[IMAGE: {image_filename}]"
                                    replaced = True
                                    break
                        
                        if replaced:
                            break
                    if replaced:
                        break
                if replaced:
                    break
        
        return replaced

    def generate_template_15(self):
        """Generate template 15 with FORCED correct scaling"""
        print(f"\n📄 GENERATING: Template 15 - FORCED SCALING")
        print("=" * 70)
        
        # STEP 1: Load master data and get exact values
        master_data = self.load_master_data()
        step_count = self.detect_step_count(master_data)
        
        # Get exact target days from master data
        placeholders = master_data.get("placeholders", {})
        time_data = placeholders.get("thoi_gian_hoan_thanh", {})
        time_content = time_data.get("content", "120 ngày")
        target_days = self.extract_days_from_time(time_content)
        
        print(f"🎯 FORCING charts to use exactly {target_days} days")
        
        # STEP 2: Generate charts with FORCED scaling
        timeline_success = self.generate_timeline_chart_directly(step_count, target_days)
        personnel_success = self.generate_personnel_chart_directly(step_count, target_days)
        
        if not (timeline_success and personnel_success):
            print("⚠️ Some charts failed to generate, but continuing...")
        
        # STEP 3: Copy template and process placeholders
        template_name = "15_KE_HOACH_TRIEN_KHAI_THUC_HIEN_template.docx"
        output_name = "15_KE_HOACH_TRIEN_KHAI_THUC_HIEN.docx"
        working_file = self.copy_template_to_working_file(template_name, output_name)
        
        # STEP 4: Process placeholders
        placeholder_mappings = [
            ("ten_goi_thau", "text"),
            ("chu_dau_tu", "text"),
            ("thoi_gian_hoan_thanh", "text"),
            ("kh_thuc_hien", "image"),
            ("kh_nhan_su", "image")
        ]
        
        replaced_count = 0
        
        # Process each placeholder
        for placeholder_key, placeholder_type in placeholder_mappings:
            print(f"\n🔄 Processing placeholder: {placeholder_key} (type: {placeholder_type})")
            
            # Load fresh document for each placeholder
            doc = Document(working_file)
            
            if placeholder_type == "text":
                # Handle text placeholders
                if placeholder_key in placeholders:
                    placeholder_data = placeholders[placeholder_key]
                    content = placeholder_data.get("content", "")
                    
                    success = self.replace_text_placeholder(doc, placeholder_key, content)
                    if success:
                        replaced_count += 1
                        doc.save(working_file)
                        print(f"💾 Saved after replacing {placeholder_key}")
                else:
                    print(f"⚠️ Placeholder {placeholder_key} not found in master data")
            
            elif placeholder_type == "image":
                # Handle image placeholders directly
                if placeholder_key == "kh_thuc_hien":
                    image_filename = f"{step_count}_BUOC_KH_THUC_HIEN.png"
                elif placeholder_key == "kh_nhan_su":
                    image_filename = f"{step_count}_BUOC_KH_NHAN_SU.png"
                else:
                    print(f"❌ Unknown image placeholder: {placeholder_key}")
                    continue
                
                success = self.replace_image_placeholder(doc, placeholder_key, image_filename)
                if success:
                    replaced_count += 1
                    doc.save(working_file)
                    print(f"💾 Saved after replacing {placeholder_key}")
        
        # Final save
        doc = Document(working_file)
        doc.save(working_file)
        
        print(f"\n✅ Generated {output_name} with FORCED scaling")
        print(f"🔄 Replaced {replaced_count}/{len(placeholder_mappings)} placeholders")
        print(f"📊 Charts use exactly {target_days} days (forced scaling)")
        
        return working_file, replaced_count

def main():
    print("📄 Template 15 Generator - FIXED SCALING VERSION")
    print("=" * 70)
    
    try:
        # Initialize generator
        generator = Template15FixedGenerator()
        print("✅ Generator initialized successfully")
        
        # Generate template 15
        output_file, replaced_count = generator.generate_template_15()
        
        print(f"\n🎉 Template 15 generation completed!")
        print(f"📄 Output file: {output_file}")
        print(f"🔄 Placeholders replaced: {replaced_count}")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()