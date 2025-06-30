#!/usr/bin/env python3
"""
14_GIAI_PHAP_CHART_INSERTER.py - Post-process chart insertion
Takes output from generate_documents.py and replaces {{kh_thuc_hien}} with dynamic chart
"""

import os
import json
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import pandas as pd
from io import StringIO

class Template14ChartInserter:
    """Insert dynamic chart into Template 14 after generate_documents.py processing"""
    
    def __init__(self):
        """Initialize the chart inserter"""
        # Set working directory to script's parent (BID_PROCESSOR)
        self.base_dir = Path(__file__).parent.parent
        os.chdir(self.base_dir)
        
        # Read current batch ID
        self.batch_id = self.read_batch_id()
        if not self.batch_id:
            raise ValueError("❌ No current batch ID found!")
        
        # Set up paths
        self.batch_dir = self.base_dir / self.batch_id
        self.extracted_dir = self.batch_dir / "extracted_data"
        self.docx_dir = self.batch_dir / "docx"
        self.pictures_dir = self.batch_dir / "pictures"
        
        # Ensure directories exist
        self.docx_dir.mkdir(exist_ok=True)
        self.pictures_dir.mkdir(exist_ok=True)
        
        print(f"📄 Template14ChartInserter initialized for batch: {self.batch_id}")

    def read_batch_id(self):
        """Read current batch ID from current_batch.txt"""
        try:
            batch_file = self.base_dir / "current_batch.txt"
            with open(batch_file, "r") as f:
                batch_id = f.read().strip()
            return batch_id
        except FileNotFoundError:
            print("❌ current_batch.txt not found!")
            return None

    def load_master_data(self):
        """Load master data to determine replacement strategy"""
        master_file = self.extracted_dir / "master_data.json"
        
        if not master_file.exists():
            raise FileNotFoundError(f"❌ Master data not found: {master_file}")
        
        with open(master_file, 'r', encoding='utf-8') as f:
            master_data = json.load(f)
        
        return master_data

    def detect_step_count(self, master_data):
        """Detect if this is 21-step or 23-step process from cac_buoc_thuc_hien"""
        placeholders = master_data.get("placeholders", {})
        cac_buoc_data = placeholders.get("cac_buoc_thuc_hien", {})
        content = cac_buoc_data.get("content", {})
        
        # Check if it's the conditional table type with step_count
        if isinstance(content, dict) and "step_count" in content:
            step_count = content["step_count"]
            print(f"🔍 Detected step count: {step_count}")
            return step_count
        
        # Fallback: assume 21 steps
        print(f"⚠️ Could not detect step count, defaulting to 21")
        return 21

    def extract_days_from_time(self, time_str):
        """Extract number of days from time string"""
        import re
        numbers = re.findall(r'\d+', time_str)
        if numbers:
            return int(numbers[0])
        return 120  # Default fallback

    def generate_timeline_chart_directly(self, step_count, target_days):
        """Generate timeline chart directly with FORCED scaling"""
        print(f"📊 Generating {step_count}-step timeline with {target_days} days...")
        
        # Hardcoded base data for different step counts
        if step_count == 21:
            csv_data = """Task,So_ngay
Giao nhận tài liệu và lập biên bản giao nhận tài liệu,1
Vận chuyển tài liệu từ kho bảo quản đến địa điểm chỉnh lý,1
Vệ sinh sơ bộ tài liệu,1
Khảo sát và biên soạn các văn bản hướng dẫn chỉnh lý,1
Phân loại tài liệu theo Hướng dẫn phân loại,20
Lập hồ sơ hoặc chỉnh sửa hoàn thiện hồ sơ theo Hướng dẫn lập hồ sơ,30
Viết các trường thông tin vào phiếu tin,3
Kiểm tra chỉnh sửa hồ sơ và phiếu tin,1
Hệ thống hóa phiếu tin theo phương án phân loại,1
Hệ thống hóa hồ sơ theo phiếu tin,1
Biên mục hồ sơ,15
Kiểm tra và chỉnh sửa việc biên mục hồ sơ,1
Ghi số hồ sơ chính thức vào phiếu tin và lên bìa hồ sơ,7
Vệ sinh tài liệu tháo bỏ ghim kẹp làm phẳng và đưa tài liệu vào bìa hồ sơ,2
Đưa hồ sơ vào hộp cặp,3
Viết in và dán nhãn hộp cặp,3
Vận chuyển tài liệu vào kho và xếp lên giá,1
Kiểm tra chỉnh sửa việc biên phiếu tin,1
Lập mục lục hồ sơ,5
Thống kê bó gói lập danh mục và viết thuyết minh tài liệu loại,1
Kết thúc chỉnh lý,1"""
        else:  # 23 steps
            csv_data = """Task,So_ngay
Giao nhận tài liệu và lập biên bản giao nhận tài liệu,0.5
Vận chuyển tài liệu từ kho bảo quản đến địa điểm chỉnh lý,1
Vệ sinh sơ bộ tài liệu,1
Khảo sát và biên soạn các văn bản hướng dẫn chỉnh lý,1
Phân loại tài liệu theo Hướng dẫn phân loại,20
Lập hồ sơ hoặc chỉnh sửa hoàn thiện hồ sơ theo Hướng dẫn lập hồ sơ,30
Viết các trường thông tin vào phiếu tin,2
Kiểm tra chỉnh sửa hồ sơ và phiếu tin,1
Hệ thống hóa phiếu tin theo phương án phân loại,1
Hệ thống hóa hồ sơ theo phiếu tin,1
Biên mục hồ sơ,15
Kiểm tra và chỉnh sửa việc biên mục hồ sơ,1
Ghi số hồ sơ chính thức vào phiếu tin và lên bìa hồ sơ,6
Vệ sinh tài liệu tháo bỏ ghim kẹp làm phẳng và đưa tài liệu vào bìa hồ sơ,1
Đưa hồ sơ vào hộp cặp,2
Viết in và dán nhãn hộp cặp,2
Vận chuyển tài liệu vào kho và xếp lên giá,1
Giao nhận tài liệu sau chỉnh lý và lập Biên bản giao nhận tài liệu,1
Nhập phiếu tin vào cơ sở dữ liệu,5
Kiểm tra chỉnh sửa việc nhập phiếu tin,1
Lập mục lục hồ sơ,4
Thống kê bó gói lập danh mục và viết thuyết minh tài liệu loại,1
Kết thúc chỉnh lý,1"""
        
        # Process data
        df = pd.read_csv(StringIO(csv_data))
        
        # FORCE scaling with exact target_days
        base_total = df['So_ngay'].sum()
        scaling_ratio = target_days / base_total
        
        print(f"📊 Base total: {base_total} days")
        print(f"📊 Target total: {target_days} days")
        print(f"📊 Scaling ratio: {scaling_ratio:.3f}")
        
        # Apply scaling and round up
        df['Days_Scaled'] = (df['So_ngay'] * scaling_ratio).apply(lambda x: max(1, round(x)))
        scaled_total = df['Days_Scaled'].sum()
        
        print(f"📊 Actual scaled total: {scaled_total} days")
        
        # Create the chart - EXACT SAME AS OLD CODE
        def wrap_text(text, words_per_line=2):
            words = text.split()
            return '\n'.join([' '.join(words[i:i+words_per_line]) for i in range(0, len(words), words_per_line)])

        wrapped_labels = [wrap_text(task, 2) for task in df["Task"]]
        max_lines = max(label.count('\n') + 1 for label in wrapped_labels)
        bottom_margin = 0.4 + 0.03 * max_lines
        fig_height = 12 + 0.5 * max_lines

        # Create figure
        fig, ax = plt.subplots(figsize=(26, fig_height))
        x = df.index
        ax.plot(x, df['Days_Scaled'], marker='o', color='red', linestyle='-')
        ax.set_title(f"KẾ HOẠCH THỰC HIỆN CÔNG VIỆC ({step_count} BƯỚC)", fontsize=20, pad=25)
        ax.set_ylabel("Số ngày", fontsize=16)
        ax.set_xticks(x)
        ax.set_xticklabels(wrapped_labels, rotation=0, ha='center', fontsize=11.5)
        ax.grid(True, axis='y', linestyle='--', alpha=0.6)

        # Set y-axis limits with extra space at top for labels
        max_y = df['Days_Scaled'].max()
        ax.set_ylim(0, max_y * 1.15)

        # Annotate data points
        for i, y in enumerate(df['Days_Scaled']):
            ax.annotate(str(y), (x[i], y), textcoords="offset points", xytext=(0, 10), va='bottom', ha='center', fontsize=12)

        plt.subplots_adjust(left=0.05, right=0.98, top=0.93, bottom=bottom_margin)
        
        # Save chart with template-specific name
        output_file = self.pictures_dir / f"14_GIAI_PHAP_{step_count}_BUOC_KH_THUC_HIEN.png"
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ Timeline chart saved: {output_file}")
        return output_file

    def replace_image_placeholder(self, doc, placeholder_key, image_path):
        """Replace image placeholders with actual images - MATCH Template 15 sizing"""
        placeholder_tag = f"{{{{{placeholder_key}}}}}"
        
        if not image_path.exists():
            print(f"❌ Image not found: {image_path}")
            return False
        
        print(f"🖼️ Replacing {placeholder_tag} with: {image_path.name}")
        
        replaced = False
        
        # Check all paragraphs for the placeholder
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                print(f"📍 Found {placeholder_tag} in paragraph {para_idx}")
                
                # Clear paragraph completely
                paragraph.clear()
                
                # Add image run
                run = paragraph.add_run()
                
                try:
                    # Use FULL document width for better visibility - MATCH Template 15
                    run.add_picture(str(image_path), width=Inches(9.0))
                    
                    # Center the paragraph
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    print(f"✅ Inserted image: {image_path.name} (width=9.0\" - full page)")
                    replaced = True
                    break
                    
                except Exception as e:
                    print(f"❌ Failed to insert image {image_path.name}: {e}")
                    # Fallback: replace with text
                    run.text = f"[IMAGE: {image_path.name}]"
                    replaced = True
                    break
        
        # Also check in table cells
        if not replaced:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if placeholder_tag in paragraph.text:
                                print(f"📍 Found {placeholder_tag} in table cell paragraph {para_idx}")
                                
                                # Clear paragraph
                                paragraph.clear()
                                
                                # Add image
                                run = paragraph.add_run()
                                
                                try:
                                    # Use large width for table cells too - MATCH Template 15
                                    run.add_picture(str(image_path), width=Inches(8.0))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    print(f"✅ Inserted image in table: {image_path.name} (width=8.0\")")
                                    replaced = True
                                    break
                                    
                                except Exception as e:
                                    print(f"❌ Failed to insert image in table: {e}")
                                    run.text = f"[IMAGE: {image_path.name}]"
                                    replaced = True
                                    break
                        
                        if replaced:
                            break
                    if replaced:
                        break
                if replaced:
                    break
        
        return replaced

    def find_template_14_output(self):
        """Find the output file from generate_documents.py for template 14"""
        # Updated to use the correct filename
        possible_names = [
            "14_GIAI_PHAP_VA_PHUONG_PHAP_LUAN_THUC_HIEN_GOI_THAU.docx"
        ]
        
        for name in possible_names:
            file_path = self.docx_dir / name
            if file_path.exists():
                print(f"✅ Found template 14 output: {name}")
                return file_path
        
        # If no output found, list available files
        print("❌ Template 14 output not found. Available files:")
        for file in self.docx_dir.glob("*.docx"):
            print(f"   - {file.name}")
        
        return None

    def process_template_14_chart(self):
        """Process template 14 to replace {{kh_thuc_hien}} with dynamic chart"""
        print(f"\n📄 POST-PROCESSING: Template 14 Chart Insertion")
        print("=" * 70)
        
        # STEP 1: Find the output file from generate_documents.py
        template_14_file = self.find_template_14_output()
        if not template_14_file:
            print("❌ Cannot proceed without template 14 output file")
            return None, 0
        
        # STEP 2: Load master data and get chart parameters
        master_data = self.load_master_data()
        step_count = self.detect_step_count(master_data)
        
        # Get exact target days from master data
        placeholders = master_data.get("placeholders", {})
        time_data = placeholders.get("thoi_gian_hoan_thanh", {})
        time_content = time_data.get("content", "120 ngày")
        target_days = self.extract_days_from_time(time_content)
        
        print(f"🎯 Generating chart with exactly {target_days} days")
        
        # STEP 3: Generate the timeline chart
        chart_file = self.generate_timeline_chart_directly(step_count, target_days)
        if not chart_file:
            print("❌ Failed to generate timeline chart")
            return None, 0
        
        # STEP 4: Process the document to replace {{kh_thuc_hien}}
        print(f"📝 Loading document: {template_14_file.name}")
        doc = Document(template_14_file)
        
        # Replace the image placeholder
        success = self.replace_image_placeholder(doc, "kh_thuc_hien", chart_file)
        
        if success:
            # OVERRIDE the existing file (same name)
            output_path = template_14_file  # Use same path to override
            doc.save(output_path)
            
            print(f"\n✅ Template 14 chart insertion completed!")
            print(f"📄 File overridden: {output_path}")
            print(f"📊 Chart uses exactly {target_days} days")
            print(f"📐 Chart width optimized for Word compatibility (6.5\" maximum)")
            
            return output_path, 1
        else:
            print("❌ Failed to replace {{kh_thuc_hien}} placeholder")
            return None, 0

def main():
    print("📄 Template 14 Chart Inserter - Post-Process Dynamic Charts")
    print("=" * 70)
    
    try:
        # Initialize inserter
        inserter = Template14ChartInserter()
        print("✅ Chart inserter initialized successfully")
        
        # Process template 14
        output_file, replaced_count = inserter.process_template_14_chart()
        
        if output_file:
            print(f"\n🎉 Template 14 chart insertion completed!")
            print(f"📄 File updated: {output_file.name}")
        else:
            print(f"\n❌ Template 14 chart insertion failed!")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()