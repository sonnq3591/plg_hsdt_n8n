#!/usr/bin/env python3
"""
Document Generator - FINAL FIXED VERSION
Selective formatting: only applies bold to the specific replaced content
"""

import os
import json
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re

class DocumentGenerator:
    """Generate final documents with SELECTIVE formatting"""
    
    def __init__(self):
        """Initialize the document generator"""
        # Set working directory to script's parent (BID_PROCESSOR)
        self.base_dir = Path(__file__).parent.parent
        os.chdir(self.base_dir)
        
        # Read current batch ID
        self.batch_id = self.read_batch_id()
        if not self.batch_id:
            raise ValueError("❌ No current batch ID found!")
        
        # Set up paths
        self.batch_dir = self.base_dir / self.batch_id
        self.extracted_dir = self.batch_dir / "extracted_data"
        self.docx_dir = self.batch_dir / "docx"
        self.templates_dir = self.base_dir / "templates"
        
        # Ensure directories exist
        self.docx_dir.mkdir(exist_ok=True)
        
        print(f"📄 DocumentGenerator initialized for batch: {self.batch_id}")

    def read_batch_id(self):
        """Read current batch ID from current_batch.txt"""
        try:
            batch_file = self.base_dir / "current_batch.txt"
            with open(batch_file, "r") as f:
                batch_id = f.read().strip()
            return batch_id
        except FileNotFoundError:
            print("❌ current_batch.txt not found!")
            return None

    def load_master_data(self):
        """Load master data to determine replacement strategy"""
        master_file = self.extracted_dir / "master_data.json"
        
        if not master_file.exists():
            raise FileNotFoundError(f"❌ Master data not found: {master_file}")
        
        with open(master_file, 'r', encoding='utf-8') as f:
            master_data = json.load(f)
        
        return master_data

    def copy_template_to_working_file(self, template_name, output_name):
        """Copy template to working file for processing"""
        template_file = self.templates_dir / template_name
        output_file = self.docx_dir / output_name
        
        if not template_file.exists():
            raise FileNotFoundError(f"❌ Template not found: {template_file}")
        
        shutil.copy2(template_file, output_file)
        print(f"✅ Copied template: {template_name} → {output_name}")
        return output_file

    def replace_simple_text_placeholder(self, doc, placeholder_key, content):
        """PRECISE FIX: Character-by-character formatting preservation"""
        placeholder_tag = f"{{{{{placeholder_key}}}}}"
        print(f"📝 Replacing {placeholder_tag} with: '{content}'")
        
        replaced = False
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            full_text = paragraph.text
            if placeholder_tag not in full_text:
                continue
            
            print(f"📍 Found {placeholder_tag} in paragraph {para_idx}")
            
            # Build character-level formatting map from original runs
            char_formatting = {}
            char_pos = 0
            
            for run in paragraph.runs:
                run_text = run.text
                run_bold = run.bold if run.bold is not None else False
                
                for char in run_text:
                    char_formatting[char_pos] = run_bold
                    char_pos += 1
            
            # Check placeholder formatting
            placeholder_start = full_text.find(placeholder_tag)
            template_bold = char_formatting.get(placeholder_start, False)
            
            print(f"   🎯 Template: Placeholder at pos {placeholder_start} should be bold = {template_bold}")
            
            # Create new text
            new_text = full_text.replace(placeholder_tag, content)
            before_text, after_text = full_text.split(placeholder_tag, 1)
            
            # Clear all runs
            for run in paragraph.runs:
                run.text = ""
            
            # Rebuild with precise character formatting
            current_run_idx = 0
            current_segment = ""
            current_bold = None
            
            def add_segment_to_run(text, bold_format):
                nonlocal current_run_idx
                if not text:
                    return
                    
                if current_run_idx < len(paragraph.runs):
                    run = paragraph.runs[current_run_idx]
                else:
                    run = paragraph.add_run()
                
                run.text = text
                run.bold = bold_format
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                current_run_idx += 1
            
            # Process "before" text with original formatting
            for i, char in enumerate(before_text):
                char_bold = char_formatting.get(i, False)
                
                if current_bold is None:
                    current_bold = char_bold
                
                if char_bold == current_bold:
                    current_segment += char
                else:
                    # Formatting changed, save current segment
                    add_segment_to_run(current_segment, current_bold)
                    current_segment = char
                    current_bold = char_bold
            
            # Save any remaining "before" segment
            if current_segment:
                add_segment_to_run(current_segment, current_bold)
                current_segment = ""
            
            # Add replacement content with template formatting
            add_segment_to_run(content, template_bold)
            
            # Process "after" text with original formatting
            after_start_pos = len(before_text) + len(placeholder_tag)
            current_bold = None
            
            for i, char in enumerate(after_text):
                original_pos = after_start_pos + i
                char_bold = char_formatting.get(original_pos, False)
                
                if current_bold is None:
                    current_bold = char_bold
                
                if char_bold == current_bold:
                    current_segment += char
                else:
                    # Formatting changed, save current segment
                    add_segment_to_run(current_segment, current_bold)
                    current_segment = char
                    current_bold = char_bold
            
            # Save any remaining "after" segment
            if current_segment:
                add_segment_to_run(current_segment, current_bold)
            
            print(f"   ✅ Applied character-precise formatting: '{content}' (bold={template_bold})")
            
            replaced = True
            break
        
        return replaced

    def replace_structured_placeholder(self, doc, placeholder_key):
        """KEEP WORKING: Your proven replacement method for structured content AND tables"""
        placeholder_tag = f"{{{{{placeholder_key}}}}}"
        source_path = self.docx_dir / f"{placeholder_key}_formatted.docx"
        
        if not source_path.exists():
            print(f"❌ Missing source doc: {source_path}")
            return False

        print(f"🔄 Using structured replacement for {placeholder_tag}")
        
        source_doc = Document(source_path)
        
        # Handle both paragraphs AND tables from source document
        source_elements = []
        for element in source_doc.element.body:
            if element.tag.endswith('}p'):  # Paragraph
                source_elements.append(('paragraph', element))
            elif element.tag.endswith('}tbl'):  # Table
                source_elements.append(('table', element))

        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                print(f"📍 Found {placeholder_tag} in paragraph {i}")
                
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)

                # Insert all elements (paragraphs AND tables) from source
                for element_type, src_element in reversed(source_elements):
                    new_element = deepcopy(src_element)
                    parent.insert(index, new_element)
                    
                    # Apply formatting for paragraphs
                    if element_type == 'paragraph':
                        try:
                            # Find the corresponding paragraph in the document
                            inserted_p = None
                            for p in doc.paragraphs:
                                if p._element == new_element:
                                    inserted_p = p
                                    break
                            
                            if inserted_p:
                                # Font preservation for paragraphs
                                for run in inserted_p.runs:
                                    run.font.size = Pt(14)
                                    run.font.name = "Times New Roman"
                        except:
                            pass  # Skip formatting errors for complex elements
                
                print(f"✅ Applied structured replacement")
                return True
        
        return False

    def replace_placeholder_hybrid(self, doc, placeholder_key, master_data):
        """FIXED: Hybrid replacement with selective formatting"""
        placeholders = master_data.get("placeholders", {})
        
        if placeholder_key not in placeholders:
            print(f"❌ Placeholder {placeholder_key} not found in master data")
            return False
        
        placeholder_data = placeholders[placeholder_key]
        content_type = placeholder_data.get("type", "unknown")
        content = placeholder_data.get("content", "")
        
        print(f"\n🎯 Processing {placeholder_key} (type: {content_type})")
        
        if content_type == "simple_text":
            return self.replace_simple_text_placeholder(doc, placeholder_key, content)
            
        elif content_type in ["structured_content", "table"]:
            return self.replace_structured_placeholder(doc, placeholder_key)
            
        else:
            print(f"❌ Unknown content type: {content_type}")
            return False

    def apply_final_formatting(self, doc):
        """Apply final formatting: 1.4 line spacing to entire document"""
        print("🎨 Applying final formatting: 1.4 line spacing...")
        
        # Apply 1.4 line spacing to all paragraphs
        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = 1.4
        
        # Apply 1.4 line spacing to table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.line_spacing = 1.4
        
        print("✅ Applied 1.4 line spacing to entire document")

    def generate_document(self, template_name, output_name, placeholders_to_replace):
        """FIXED: Generate document with selective formatting"""
        print(f"\n📄 GENERATING: {output_name}")
        print("=" * 60)
        
        # Load master data
        master_data = self.load_master_data()
        
        # Copy template to working file
        working_file = self.copy_template_to_working_file(template_name, output_name)
        
        # Process placeholders ONE BY ONE with fresh document loads
        replaced_count = 0
        
        for placeholder_key in placeholders_to_replace:
            print(f"\n🔄 Processing placeholder: {placeholder_key}")
            
            # Load fresh document for each placeholder to avoid contamination
            doc = Document(working_file)
            
            success = self.replace_placeholder_hybrid(doc, placeholder_key, master_data)
            
            if success:
                replaced_count += 1
                # Save immediately after each successful replacement
                doc.save(working_file)
                print(f"💾 Saved document after replacing {placeholder_key}")
            else:
                print(f"❌ Failed to replace {placeholder_key}")
        
        # Final formatting and save
        print(f"\n🎨 Applying final formatting...")
        doc = Document(working_file)
        self.apply_final_formatting(doc)
        doc.save(working_file)
        
        print(f"✅ Generated {output_name}")
        print(f"🔄 Replaced {replaced_count}/{len(placeholders_to_replace)} placeholders")
        
        return working_file, replaced_count

    def load_formatting_summary(self):
        """Load formatting summary to know what content is available"""
        summary_file = self.docx_dir / "formatting_summary.json"
        
        if not summary_file.exists():
            raise FileNotFoundError(f"❌ Formatting summary not found: {summary_file}")
        
        with open(summary_file, 'r', encoding='utf-8') as f:
            summary = json.load(f)
        
        return summary

    def generate_all_documents(self):
        """Generate all documents with selective formatting"""
        print("\n📄 GENERATING: All Documents with Selective Formatting")
        print("=" * 60)
        
        # Load formatting summary
        try:
            summary = self.load_formatting_summary()
        except FileNotFoundError as e:
            print(e)
            return False
        
        available_placeholders = summary.get("formatted_placeholders", [])
        
        if not available_placeholders:
            print("❌ No formatted content available")
            return False
        
        print(f"📊 Available placeholders: {available_placeholders}")
        
        # Document generation rules - ALL 4 TEMPLATES
        documents_to_generate = [
            {
                "template": "02_MUC_DO_HIEU_BIET_template.docx",
                "output": "02_MUC_DO_HIEU_BIET.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "04_CAM_KET_DAP_UNG_YEU_CAU_CHUONG_V_template.docx",
                "output": "04_CAM_KET_DAP_UNG_YEU_CAU_CHUONG_V.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "05_CAM_KET_DAP_UNG_VPP_template.docx",
                "output": "05_CAM_KET_DAP_UNG_VPP.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "06_TINH_BAO_MAT_template.docx",
                "output": "06_TINH_BAO_MAT.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "08_CAM_KET_THUC_PM_template.docx",
                "output": "08_CAM_KET_THUC_PM.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "10_QUY_DINH_AP_DUNG_template.docx",
                "output": "10_QUY_DINH_AP_DUNG.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "11_CAM_KET_THUC_HIEN_GOI_THAU_template.docx", 
                "output": "11_CAM_KET_THUC_HIEN_GOI_THAU.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "12_CAM_KET_BAO_HANH_XU_LY_SU_CO_template.docx",
                "output": "12_CAM_KET_BAO_HANH_XU_LY_SU_CO.docx", 
                "placeholders": available_placeholders
            },
            {
                "template": "14_GIAI_PHAP_VA_PHUONG_PHAP_LUAN_THUC_HIEN_GOI_THAU_template.docx",
                "output": "14_GIAI_PHAP_VA_PHUONG_PHAP_LUAN_THUC_HIEN_GOI_THAU.docx", 
                "placeholders": available_placeholders
            }
        ]
        
        generated_documents = []
        
        # Generate each document
        for doc_config in documents_to_generate:
            template_name = doc_config["template"]
            output_name = doc_config["output"]
            placeholders = doc_config["placeholders"]
            
            # Filter to only use available placeholders
            available_placeholders_for_doc = [p for p in placeholders if p in available_placeholders]
            
            if not available_placeholders_for_doc:
                print(f"⚠️ Skipping {output_name} - no available placeholders")
                continue
            
            try:
                output_file, replaced_count = self.generate_document(
                    template_name, 
                    output_name, 
                    available_placeholders_for_doc
                )
                
                generated_documents.append({
                    "file": str(output_file),
                    "template": template_name,
                    "placeholders_replaced": replaced_count,
                    "total_placeholders": len(available_placeholders_for_doc)
                })
                
            except Exception as e:
                print(f"❌ Failed to generate {output_name}: {str(e)}")
                import traceback
                traceback.print_exc()
                continue
        
        # Save generation summary
        generation_summary = {
            "batch_id": self.batch_id,
            "generated_documents": generated_documents,
            "total_documents": len(generated_documents),
            "available_placeholders": available_placeholders
        }
        
        summary_file = self.docx_dir / "generation_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(generation_summary, f, ensure_ascii=False, indent=2)
        
        print(f"\n✅ SUCCESS: Generated {len(generated_documents)} documents!")
        for doc in generated_documents:
            print(f"📄 {Path(doc['file']).name}: {doc['placeholders_replaced']}/{doc['total_placeholders']} placeholders")
        
        return len(generated_documents) > 0

def main():
    print("📄 Document Generator - FINAL FIXED VERSION")
    print("=" * 70)
    
    try:
        # Initialize generator
        generator = DocumentGenerator()
        print("✅ Generator initialized successfully")
        
        # Generate all documents
        success = generator.generate_all_documents()
        
        if success:
            print("\n🎉 Document generation completed successfully!")
        else:
            print("\n❌ Document generation failed!")
            
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()