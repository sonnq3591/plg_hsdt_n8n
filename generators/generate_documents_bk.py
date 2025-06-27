#!/usr/bin/env python3
"""
Document Generator - FIXED with Hybrid Replacement Logic
Combines run-level replacement for simple text with your proven paragraph replacement
"""

import os
import json
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re

class DocumentGenerator:
    """Generate final documents with HYBRID replacement logic"""
    
    def __init__(self):
        """Initialize the document generator"""
        # Set working directory to script's parent (BID_PROCESSOR)
        self.base_dir = Path(__file__).parent.parent
        os.chdir(self.base_dir)
        
        # Read current batch ID
        self.batch_id = self.read_batch_id()
        if not self.batch_id:
            raise ValueError("❌ No current batch ID found!")
        
        # Set up paths
        self.batch_dir = self.base_dir / self.batch_id
        self.extracted_dir = self.batch_dir / "extracted_data"
        self.docx_dir = self.batch_dir / "docx"
        self.templates_dir = self.base_dir / "templates"
        
        # Ensure directories exist
        self.docx_dir.mkdir(exist_ok=True)
        
        print(f"📄 DocumentGenerator initialized for batch: {self.batch_id}")

    def read_batch_id(self):
        """Read current batch ID from current_batch.txt"""
        try:
            batch_file = self.base_dir / "current_batch.txt"
            with open(batch_file, "r") as f:
                batch_id = f.read().strip()
            return batch_id
        except FileNotFoundError:
            print("❌ current_batch.txt not found!")
            return None

    def load_master_data(self):
        """Load master data to determine replacement strategy"""
        master_file = self.extracted_dir / "master_data.json"
        
        if not master_file.exists():
            raise FileNotFoundError(f"❌ Master data not found: {master_file}")
        
        with open(master_file, 'r', encoding='utf-8') as f:
            master_data = json.load(f)
        
        return master_data

    def copy_template_to_working_file(self, template_name, output_name):
        """Copy template to working file for processing"""
        template_file = self.templates_dir / template_name
        output_file = self.docx_dir / output_name
        
        if not template_file.exists():
            raise FileNotFoundError(f"❌ Template not found: {template_file}")
        
        shutil.copy2(template_file, output_file)
        print(f"✅ Copied template: {template_name} → {output_name}")
        return output_file

    def replace_simple_text_placeholder(self, doc, placeholder_key, content):
        """Use DOCUMENT 02 PROVEN METHOD: Create formatted DOCX and use paragraph replacement"""
        placeholder_tag = f"{{{{{placeholder_key}}}}}"
        print(f"📝 Using Document 02 proven method for {placeholder_tag}")
        
        # Create a temporary formatted DOCX file using Document 02's approach
        temp_doc = Document()
        style = temp_doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        
        # Add the content as a simple paragraph
        para = temp_doc.add_paragraph(content)
        
        # Apply Document 02's proven formatting
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        
        # Ensure font formatting
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        # Save temporary file
        temp_path = self.docx_dir / f"{placeholder_key}_formatted.docx"
        temp_doc.save(temp_path)
        print(f"✅ Created formatted file: {temp_path}")
        
        # Now use the PROVEN paragraph replacement method
        return self.replace_structured_placeholder(doc, placeholder_key)

    def replace_structured_placeholder(self, doc, placeholder_key):
        """YOUR PROVEN replacement method for structured content AND tables"""
        placeholder_tag = f"{{{{{placeholder_key}}}}}"
        source_path = self.docx_dir / f"{placeholder_key}_formatted.docx"
        
        if not source_path.exists():
            print(f"❌ Missing source doc: {source_path}")
            return False

        print(f"🔄 Using YOUR PROVEN method for {placeholder_tag}")
        
        source_doc = Document(source_path)
        
        # Handle both paragraphs AND tables from source document
        source_elements = []
        for element in source_doc.element.body:
            if element.tag.endswith('}p'):  # Paragraph
                source_elements.append(('paragraph', element))
            elif element.tag.endswith('}tbl'):  # Table
                source_elements.append(('table', element))

        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_tag in paragraph.text:
                print(f"📍 Found {placeholder_tag} in paragraph {i}")
                
                p_element = paragraph._element
                parent = p_element.getparent()
                index = parent.index(p_element)
                parent.remove(p_element)

                # Insert all elements (paragraphs AND tables) from source
                for element_type, src_element in reversed(source_elements):
                    new_element = deepcopy(src_element)
                    parent.insert(index, new_element)
                    
                    # Apply formatting for paragraphs
                    if element_type == 'paragraph':
                        try:
                            # Find the corresponding paragraph in the document
                            inserted_p = None
                            for p in doc.paragraphs:
                                if p._element == new_element:
                                    inserted_p = p
                                    break
                            
                            if inserted_p:
                                # YOUR PROVEN FONT PRESERVATION for paragraphs
                                for run in inserted_p.runs:
                                    run.font.size = Pt(14)
                                    run.font.name = "Times New Roman"
                        except:
                            pass  # Skip formatting errors for complex elements
                
                print(f"✅ Applied YOUR PROVEN replacement method with table support")
                return True
        
        return False

    def replace_placeholder_hybrid(self, doc, placeholder_key, master_data):
        """HYBRID replacement: Choose method based on content type"""
        placeholders = master_data.get("placeholders", {})
        
        if placeholder_key not in placeholders:
            print(f"❌ Placeholder {placeholder_key} not found in master data")
            return False
        
        placeholder_data = placeholders[placeholder_key]
        content_type = placeholder_data.get("type", "unknown")
        
        print(f"🎯 Processing {placeholder_key} as type: {content_type}")
        
        if content_type == "simple_text":
            # Use run-level replacement for simple text
            content = placeholder_data.get("content", "")
            return self.replace_simple_text_placeholder(doc, placeholder_key, content)
            
        elif content_type in ["structured_content", "table"]:
            # Use YOUR PROVEN paragraph replacement for complex content
            return self.replace_structured_placeholder(doc, placeholder_key)
            
        else:
            print(f"❌ Unknown content type: {content_type}")
            return False

    def apply_final_formatting(self, doc):
        """Apply final formatting: 1.4 line spacing to entire document"""
        print("🎨 Applying final formatting: 1.4 line spacing...")
        
        from docx.shared import Inches
        
        # Apply 1.4 line spacing to all paragraphs
        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = 1.4
        
        # Apply 1.4 line spacing to table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.line_spacing = 1.4
        
        print("✅ Applied 1.4 line spacing to entire document")

    def generate_document(self, template_name, output_name, placeholders_to_replace):
        """Generate document with HYBRID replacement logic"""
        print(f"\n📄 GENERATING: {output_name}")
        print("=" * 60)
        
        # Load master data to determine replacement strategies
        master_data = self.load_master_data()
        
        # Copy template to working file
        working_file = self.copy_template_to_working_file(template_name, output_name)
        
        # Load the document
        doc = Document(working_file)
        
        # Replace each placeholder using HYBRID logic
        replaced_count = 0
        for placeholder_key in placeholders_to_replace:
            success = self.replace_placeholder_hybrid(doc, placeholder_key, master_data)
            if success:
                replaced_count += 1
        
        # Save the final document with 1.4 line spacing
        self.apply_final_formatting(doc)
        doc.save(working_file)
        
        print(f"✅ Generated {output_name}")
        print(f"🔄 Replaced {replaced_count}/{len(placeholders_to_replace)} placeholders")
        
        return working_file, replaced_count

    def load_formatting_summary(self):
        """Load formatting summary to know what content is available"""
        summary_file = self.docx_dir / "formatting_summary.json"
        
        if not summary_file.exists():
            raise FileNotFoundError(f"❌ Formatting summary not found: {summary_file}")
        
        with open(summary_file, 'r', encoding='utf-8') as f:
            summary = json.load(f)
        
        return summary

    def generate_all_documents(self):
        """Generate all documents with hybrid replacement"""
        print("\n📄 GENERATING: All Documents with Hybrid Logic")
        print("=" * 60)
        
        # Load formatting summary
        try:
            summary = self.load_formatting_summary()
        except FileNotFoundError as e:
            print(e)
            return False
        
        available_placeholders = summary.get("formatted_placeholders", [])
        
        if not available_placeholders:
            print("❌ No formatted content available")
            return False
        
        print(f"📊 Available placeholders: {available_placeholders}")
        
        # Document generation rules - ALL 4 TEMPLATES
        documents_to_generate = [
            {
                "template": "02_MUC_DO_HIEU_BIET_template.docx",
                "output": "02_MUC_DO_HIEU_BIET_output.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "04_CAM_KET_DAP_UNG_YEU_CAU_CHUONG_V_template.docx",
                "output": "04_CAM_KET_DAP_UNG_YEU_CAU_CHUONG_V_output.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "11_CAM_KET_THUC_HIEN_GOI_THAU_template.docx", 
                "output": "11_CAM_KET_THUC_HIEN_GOI_THAU_output.docx",
                "placeholders": available_placeholders
            },
            {
                "template": "12_CAM_KET_BAO_HANH_XU_LY_SU_CO_template.docx",
                "output": "12_CAM_KET_BAO_HANH_XU_LY_SU_CO_output.docx", 
                "placeholders": available_placeholders
            }
        ]
        
        generated_documents = []
        
        # Generate each document
        for doc_config in documents_to_generate:
            template_name = doc_config["template"]
            output_name = doc_config["output"]
            placeholders = doc_config["placeholders"]
            
            # Filter to only use available placeholders
            available_placeholders_for_doc = [p for p in placeholders if p in available_placeholders]
            
            if not available_placeholders_for_doc:
                print(f"⚠️ Skipping {output_name} - no available placeholders")
                continue
            
            try:
                output_file, replaced_count = self.generate_document(
                    template_name, 
                    output_name, 
                    available_placeholders_for_doc
                )
                
                generated_documents.append({
                    "file": str(output_file),
                    "template": template_name,
                    "placeholders_replaced": replaced_count,
                    "total_placeholders": len(available_placeholders_for_doc)
                })
                
            except Exception as e:
                print(f"❌ Failed to generate {output_name}: {str(e)}")
                import traceback
                traceback.print_exc()
                continue
        
        # Save generation summary
        generation_summary = {
            "batch_id": self.batch_id,
            "generated_documents": generated_documents,
            "total_documents": len(generated_documents),
            "available_placeholders": available_placeholders
        }
        
        summary_file = self.docx_dir / "generation_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(generation_summary, f, ensure_ascii=False, indent=2)
        
        print(f"\n✅ SUCCESS: Generated {len(generated_documents)} documents!")
        for doc in generated_documents:
            print(f"📄 {Path(doc['file']).name}: {doc['placeholders_replaced']}/{doc['total_placeholders']} placeholders")
        
        return len(generated_documents) > 0

def main():
    print("📄 Document Generator - FIXED with Hybrid Replacement Logic")
    print("=" * 70)
    
    try:
        # Initialize generator
        generator = DocumentGenerator()
        print("✅ Generator initialized successfully")
        
        # Generate all documents
        success = generator.generate_all_documents()
        
        if success:
            print("\n🎉 Document generation completed successfully!")
        else:
            print("\n❌ Document generation failed!")
            
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()