#!/usr/bin/env python3
"""
Structured Content Formatter
Uses YOUR PROVEN markdown_to_docx method for can_cu_phap_ly and muc_dich_cong_viec
"""

import os
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class StructuredContentFormatter:
    def __init__(self, batch_id):
        self.batch_id = batch_id
        self.base_dir = Path(__file__).parent.parent
        self.batch_dir = self.base_dir / batch_id
        self.extracted_dir = self.batch_dir / "extracted_data"
        self.docx_dir = self.batch_dir / "docx"
        
        print(f"📋 StructuredContentFormatter initialized for batch: {batch_id}")

    def format_structured_content(self, placeholder_key, structured_data):
        """Format structured content using YOUR PROVEN markdown_to_docx method"""
        print(f"📋 Formatting {placeholder_key} as structured content")
        
        if not structured_data:
            print("⚠️ Empty structured content")
            return None
        
        # Convert structured data back to markdown format
        markdown_lines = []
        
        for element in structured_data:
            element_type = element.get("type", "paragraph")
            text = element.get("text", "")
            
            if element_type == "heading":
                # Add as **bold** heading
                markdown_lines.append(f"**{text}**")
                
                # Add bullet items for this heading
                items = element.get("items", [])
                for item in items:
                    markdown_lines.append(f"- {item}")
                    
            elif element_type == "bullet_item":
                # Standalone bullet item
                markdown_lines.append(f"- {text}")
                    
            elif element_type == "paragraph":
                # Regular paragraph
                markdown_lines.append(text)
        
        # Join to create markdown content
        markdown_content = '\n'.join(markdown_lines)
        print(f"📝 Generated markdown: {len(markdown_lines)} lines")
        
        # Create document with YOUR PROVEN method
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # YOUR PROVEN markdown_to_docx logic - EXACTLY as you wrote it
        for line in markdown_content.splitlines():
            line = line.strip()
            if not line:
                continue
                
            if re.match(r'^\*\*(.+)\*\*$', line):
                # YOUR PROVEN bold heading logic
                text = re.sub(r'\*\*(.+)\*\*', r'\1', line)
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
            elif line.startswith('- '):
                # YOUR PROVEN bullet point logic
                para = doc.add_paragraph(f'- {line[2:].strip()}')
            else:
                # YOUR PROVEN regular paragraph logic
                para = doc.add_paragraph(line)
            
            # YOUR PROVEN formatting - EXACTLY as you do it
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(6)
        
        # Save
        output_path = self.docx_dir / f"{placeholder_key}_formatted.docx"
        doc.save(output_path)
        print(f"✅ Saved structured content: {output_path}")
        
        return output_path

def main():
    # Test standalone
    if len(sys.argv) > 1:
        batch_id = sys.argv[1]
        formatter = StructuredContentFormatter(batch_id)
        # Test with sample structured data
        sample_data = [
            {"type": "paragraph", "text": "Sample paragraph"},
            {"type": "heading", "text": "Sample Heading", "items": ["Item 1", "Item 2"]},
            {"type": "bullet_item", "text": "Standalone bullet"}
        ]
        formatter.format_structured_content("test_structured", sample_data)

if __name__ == "__main__":
    import sys
    main()