#!/usr/bin/env python3
"""
Table Formatter
Formats table content with your proven table logic
"""

import os
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime

class TableFormatter:
    def __init__(self, batch_id):
        self.batch_id = batch_id
        self.base_dir = Path(__file__).parent.parent
        self.batch_dir = self.base_dir / batch_id
        self.extracted_dir = self.batch_dir / "extracted_data"
        self.docx_dir = self.batch_dir / "docx"
        
        print(f"📊 TableFormatter initialized for batch: {batch_id}")

    def format_table(self, placeholder_key, table_data):
        """Format table using your proven table logic"""
        print(f"📊 Formatting {placeholder_key} as table")
        
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        
        if not headers or not rows:
            print("⚠️ Empty table data")
            return None
        
        # Create document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        
        # Create table
        table = doc.add_table(rows=0, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Calculate column widths (your proven logic)
        max_lengths = [0] * len(headers)
        
        # Check headers
        for i, header in enumerate(headers):
            max_lengths[i] = max(max_lengths[i], len(str(header)))
        
        # Check data rows
        for row in rows:
            for i in range(min(len(row), len(headers))):
                max_lengths[i] = max(max_lengths[i], len(str(row[i])))
        
        # Set column widths
        total_width = 7.0
        min_width = 0.6
        max_width = 2.0
        
        total_chars = sum(max_lengths) if sum(max_lengths) > 0 else len(headers)
        column_widths = []
        
        for length in max_lengths:
            if total_chars > 0:
                width = (length / total_chars) * total_width
            else:
                width = total_width / len(headers)
            width = max(min_width, min(width, max_width))
            column_widths.append(width)
        
        # Add header row
        header_row = table.add_row().cells
        for i, header_text in enumerate(headers):
            cell = header_row[i]
            cell.width = Inches(column_widths[i])
            
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(str(header_text).strip())
            
            # Header formatting
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Add data rows
        for row_data in rows:
            data_row = table.add_row().cells
            
            for i in range(len(headers)):
                cell_text = str(row_data[i]) if i < len(row_data) else ""
                cell = data_row[i]
                cell.width = Inches(column_widths[i])
                
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(cell_text.strip())
                
                # Data cell formatting
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                
                # Smart alignment
                if i == 0 or len(cell_text.strip()) < 10:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Save
        output_path = self.docx_dir / f"{placeholder_key}_formatted.docx"
        doc.save(output_path)
        print(f"✅ Saved table: {output_path}")
        
        return output_path

def main():
    # Test standalone
    if len(sys.argv) > 1:
        batch_id = sys.argv[1]
        formatter = TableFormatter(batch_id)
        # Test with sample data
        sample_data = {
            "headers": ["STT", "Name", "Value"],
            "rows": [["1", "Test", "123"], ["2", "Test2", "456"]]
        }
        formatter.format_table("test_table", sample_data)

if __name__ == "__main__":
    import sys
    main()