#!/usr/bin/env python3
"""
Conditional Table Formatter - NEW Specialized Formatter
Handles conditional table content based on step count (21_BUOC.docx or 23_BUOC.docx)
"""

import os
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from datetime import datetime

class ConditionalTableFormatter:
    def __init__(self, batch_id):
        self.batch_id = batch_id
        self.base_dir = Path(__file__).parent.parent
        self.batch_dir = self.base_dir / batch_id
        self.extracted_dir = self.batch_dir / "extracted_data"
        self.docx_dir = self.batch_dir / "docx"
        self.templates_dir = self.base_dir / "templates"
        
        print(f"📊 ConditionalTableFormatter initialized for batch: {batch_id}")

    def format_conditional_table(self, placeholder_key, content_data):
        """Format conditional table using EXACT method from processor_cac_buoc_premade.py"""
        print(f"📊 Formatting {placeholder_key} as conditional table")
        
        # Extract conditional logic data
        step_count = content_data.get("step_count", 21)
        source_file = content_data.get("source_file", "21_BUOC.docx")
        
        print(f"🔢 Step count: {step_count}")
        print(f"📄 Source file: {source_file}")
        
        # Load the appropriate template file
        source_path = self.templates_dir / source_file
        if not source_path.exists():
            print(f"❌ Source template not found: {source_path}")
            return None
        
        print(f"📖 Loading source template: {source_path}")
        
        try:
            # Load source document
            source_doc = Document(source_path)
            
            # Create new document with standard formatting
            output_doc = Document()
            style = output_doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            
            # EXACT METHOD from processor_cac_buoc_premade.py
            print(f"📊 Source document has {len(source_doc.paragraphs)} paragraphs")
            print(f"📊 Source document has {len(source_doc.tables)} tables")
            
            # Get all content (paragraphs and tables) - EXACT COPY
            all_elements = []
            
            # Add paragraphs - EXACT COPY
            for para in source_doc.paragraphs:
                if para.text.strip():
                    all_elements.append(('paragraph', para))
            
            # Add tables - EXACT COPY
            for table in source_doc.tables:
                all_elements.append(('table', table))
            
            print(f"📊 Total elements to copy: {len(all_elements)}")
            
            # Get the document body for XML element insertion
            doc_body = output_doc.element.body
            
            # Insert all elements using XML element copying - EXACT METHOD
            for element_type, element in all_elements:
                if element_type == 'paragraph':
                    from copy import deepcopy
                    new_p = deepcopy(element._element)  # XML element level!
                    doc_body.append(new_p)
                elif element_type == 'table':
                    from copy import deepcopy
                    new_t = deepcopy(element._element)  # XML element level!
                    doc_body.append(new_t)
            
            print(f"✅ Copied all {len(all_elements)} elements using XML element method")
            
            # Apply justified alignment and fix first column letters only
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            print(f"🎨 Applying justified alignment and fixing first column letters...")
            
            # Fix ONLY first column letters in tables
            for table in output_doc.tables:
                for row in table.rows:
                    first_cell = row.cells[0]
                    cell_text = first_cell.text.strip()
                    
                    # If it's a letter and doesn't have ")", add ")"
                    if (len(cell_text) == 1 and 
                        cell_text.isalpha() and 
                        ')' not in cell_text):
                        print(f"🔧 Fixing: '{cell_text}' → '{cell_text})'")
                        first_cell.text = f"{cell_text})"
                    
                    # If first column has letter with ")", make it centered and italic
                    if (len(cell_text) <= 3 and 
                        cell_text.replace(')', '').isalpha()):
                        print(f"🎨 Styling: '{cell_text}' → centered + italic")
                        
                        # Clear and rebuild with formatting
                        first_cell.text = ""
                        para = first_cell.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = para.add_run(cell_text if ')' in cell_text else f"{cell_text})")
                        run.italic = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
            
            # Apply justified alignment to paragraphs
            for para in output_doc.paragraphs:
                if para.text.strip():
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            print(f"✅ Done")
            
            # Save formatted document
            output_path = self.docx_dir / f"{placeholder_key}_formatted.docx"
            output_doc.save(output_path)
            print(f"✅ Saved conditional table: {output_path}")
            
            return output_path
            
        except Exception as e:
            print(f"❌ Error formatting conditional table: {str(e)}")
            return None

def main():
    # Test standalone
    if len(sys.argv) > 1:
        batch_id = sys.argv[1]
        formatter = ConditionalTableFormatter(batch_id)
        
        # Test with sample data
        sample_data = {
            "step_count": 21,
            "source_file": "21_BUOC.docx",
            "conditional_logic": True
        }
        formatter.format_conditional_table("test_conditional", sample_data)

if __name__ == "__main__":
    import sys
    main()